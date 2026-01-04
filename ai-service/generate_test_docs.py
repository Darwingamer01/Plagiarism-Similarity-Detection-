
import os
from docx import Document
import random

# Ensure output directory exists
OUTPUT_DIR = "../test_docs_v2"
if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

def create_docx(filename, content):
    doc = Document()
    doc.add_heading(filename.replace('_', ' ').replace('.docx', ''), 0)
    for paragraph in content.split('\n\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    path = os.path.join(OUTPUT_DIR, filename)
    doc.save(path)
    print(f"Created: {path}")

# --- Content Generation Helpers ---

def get_base_text():
    # Approx 600 words base, we will repeat/expand to make it "Big"
    return """The Transformative Role of Artificial Intelligence in Modern Healthcare Systems

Artificial Intelligence (AI) is rapidly transforming the landscape of modern healthcare, offering unprecedented opportunities to improve patient outcomes, enhance diagnostic accuracy, and streamline administrative processes. From predictive analytics to robotic surgery, AI technologies are being integrated into various facets of medical practice, fundamentally altering how care is delivered and experienced. This essay explores the multifaceted role of AI in healthcare, examining its benefits, challenges, and future potential.

One of the most significant contributions of AI in healthcare is in the realm of medical imaging and diagnostics. Machine learning algorithms, particularly deep learning models, have demonstrated remarkable proficiency in analyzing complex medical images such as X-rays, MRI scans, and CT scans. These algorithms can detect anomalies like tumors, fractures, and retinal diseases with accuracy often comparable to, or even exceeding, that of human radiologists. For instance, AI-powered tools are now being used to screen for diabetic retinopathy, a leading cause of blindness, allowing for earlier detection and intervention. By augmenting the capabilities of clinicians, AI ensures more timely and accurate diagnoses, which are critical for effective treatment planning.

Beyond diagnostics, AI is revolutionizing personalized medicine. Traditional medical treatments often follow a "one-size-fits-all" approach, which may not be effective for every patient. AI enables the analysis of vast datasets, including genomic information, electronic health records, and lifestyle factors, to tailor treatments to individual patients. This precision medicine approach allows oncologists to identify specific genetic mutations driving a patient's cancer and recommend targeted therapies that are more likely to be effective. Similarly, AI models can predict how patients will respond to certain drugs, minimizing adverse effects and optimizing dosage.

In the surgical theater, AI-driven robotics are enhancing the precision and safety of procedures. Robotic surgical systems, guided by AI, allow surgeons to perform complex, minimally invasive operations with greater dexterity and control than is possible with the human hand alone. These systems can filter out hand tremors and provide real-time feedback during surgery, reducing the risk of complications and shortening recovery times. Furthermore, preoperative planning using AI algorithms helps surgeons visualize the patient's anatomy in 3D, allowing for meticulous preparation before the first incision is made.

Administrative efficiency is another area where AI is making a substantial impact. Healthcare systems are burdened with enormous amounts of paperwork, billing, and scheduling tasks that consume valuable time and resources. Natural Language Processing (NLP) technologies can automate clinical documentation by transcribing doctor-patient interactions and populating electronic health records. AI chatbots and virtual assistants handle appointment scheduling and answer routine patient queries, freeing up staff to focus on more complex tasks. This operational efficiency not only reduces costs but also alleviates burnout among healthcare professionals, allowing them to dedicate more time to direct patient care.

However, the integration of AI in healthcare is not without challenges. Data privacy and security are paramount concerns, as the training of AI models requires access to massive amounts of sensitive patient data. Ensuring that this data is anonymized and protected from breaches is a critical ethical and legal obligation. Additionally, the potential for algorithmic bias poses a significant risk. If AI models are trained on data that is not representative of diverse populations, they may produce biased results that exacerbate existing healthcare disparities. It is essential to develop and validate AI tools on diverse datasets to ensure equitable care for all patients.

In conclusion, Artificial Intelligence stands at the forefront of a healthcare revolution. By enhancing diagnostic precision, enabling personalized treatments, improving surgical outcomes, and optimizing administrative workflows, AI holds the promise of a more efficient and effective healthcare system. However, realizing this potential requires a balanced approach that addresses ethical concerns, ensures data security, and actively mitigates bias. As technology continues to evolve, the collaboration between human intelligence and artificial intelligence will undoubtedly define the future of medicine, ultimately leading to healthier lives for people around the world."""

def get_base_content_large():
    base = get_base_text()
    # Repeat key sections to increase size to "Big Content" (approx 2000-3000 words)
    expanded = base + "\n\n" + \
               "### Extended Analysis on Diagnostics\n" + \
               (base.split('\n\n')[1] + " ")*3 + "\n\n" + \
               "### Deep Dive into Robotic Surgery\n" + \
               (base.split('\n\n')[3] + " ")*3 + "\n\n" + \
               "### Administrative Challenges in Depth\n" + \
               (base.split('\n\n')[4] + " ")*3 + "\n\n" + \
               base  # Repeat full cycle
    return expanded

def mix_content(content1, content2, ratio):
    # Rudimentary mixing: take first ratio% of content1 and remaining (1-ratio)% of content2
    words1 = content1.split()
    words2 = content2.split()
    
    cutoff1 = int(len(words1) * ratio)
    cutoff2 = int(len(words2) * (1 - ratio))
    
    mixed_words = words1[:cutoff1] + words2[:cutoff2]
    return ' '.join(mixed_words)

def generate_paraphrased(text):
    # A simple programmatic "paraphrase" by replacing common words
    # In a real scenario, use an LLM. Here we simulates structural change.
    replacements = {
        "AI": "Artificial Intelligence",
        "healthcare": "medical care",
        "important": "crucial",
        "transforming": "changing",
        "rapidly": "quickly",
        "modern": "contemporary",
        "provides": "offers",
        "challenges": "difficulties",
        "benefits": "advantages",
        "essay": "article",
        "significant": "major",
        "realm": "field",
        "proficiency": "skill",
        "complex": "complicated",
        "anomalies": "abnormalities",
        "clinicians": "doctors",
        "augmenting": "increasing",
        "ensure": "make sure",
        "timely": "prompt",
        "accurate": "precise",
        "tailor": "customize",
        "vast": "huge",
        "minimize": "reduce",
        "adverse": "bad",
        "dexterity": "skill",
        "tremors": "shakes",
        "meticulous": "careful",
        "burdened": "overloaded",
        "enormous": "huge",
        "consume": "use up",
        "alleviates": "eases",
        "burnout": "exhaustion",
        "paramount": "top",
        "breaches": "leaks",
        "bias": "prejudice",
        "exacerbate": "worsen",
        "disparities": "inequalities",
        "equitable": "fair",
        "forefront": "leading edge",
        "mitigates": "reduces"
    }
    
    words = text.split()
    new_words = []
    for word in words:
        clean_word = word.strip(".,!?\"")
        if clean_word in replacements and random.random() > 0.5:
             new_words.append(replacements[clean_word])
        else:
             new_words.append(word)
    
    # Shuffle sentence structures slightly (simple simulation)
    sentences = ' '.join(new_words).split('. ')
    # Don't shuffle too much to keep coherence for "structure change" check
    return '. '.join(sentences)

def get_irrelevant_content():
    return """The Universe is all of space and time and their contents, including planets, stars, galaxies, and all other forms of matter and energy. The Big Bang theory is the prevailing cosmological description of the development of the universe. According to this theory, space and time emerged together 13.787±0.020 billion years ago, and the universe has been expanding ever since. While the spatial size of the entire universe is unknown, it is possible to measure the size of the observable universe, which is currently estimated to be 93 billion light-years in diameter.

The earliest cosmological models of the universe were developed by ancient Greek and Indian philosophers and were geocentric, placing Earth at the center. Over the centuries, more precise astronomical observations led Nicolaus Copernicus to develop the heliocentric model with the Sun at the center of the Solar System. In developing the law of universal gravitation, Isaac Newton built upon Copernicus's work as well as observations by Tycho Brahe and Johannes Kepler's laws of planetary motion.

Further observational improvements led to the realization that the Sun is one of the hundreds of billions of stars in the Milky Way, which is one of at least two trillion galaxies in the universe. Many of the stars in our galaxy have planets. At the largest scale, galaxies are distributed uniformly and the same in all directions, meaning that the universe has neither an edge nor a center. At smaller scales, galaxies are distributed in filaments and voids in a vast cosmic web.

The detailed chemical composition of the universe is dominated by hydrogen and helium atoms. These are the most abundant elements, produced in the Big Bang. Heavier elements, such as carbon, nitrogen, and oxygen, are produced in the cores of stars through nuclear fusion. When these stars die, they release these elements into space, enriching the interstellar medium and allowing for the formation of new stars and planets.

Black holes are among the most intriguing objects in the universe. They are regions of spacetime where gravity is so strong that nothing—no particles or even electromagnetic radiation such as light—can escape from it. The theory of general relativity predicts that a sufficiently compact mass can deform spacetime to form a black hole. In many ways, black holes act as the ultimate recyclers of the universe, consuming matter and energy and potentially influencing the evolution of galaxies.

Dark matter and dark energy are two of the greatest mysteries in modern cosmology. Dark matter is a form of matter that is thought to account for approximately 85% of the matter in the universe. It does not interact with electromagnetic radiation, making it invisible to telescopes. Its existence is inferred from its gravitational effects on visible matter, such as the rotation curves of galaxies. Dark energy, on the other hand, is a hypothetical form of energy that permeates all of space and tends to accelerate the expansion of the universe. Together, these unknown components make up about 95% of the total mass-energy content of the universe.""" * 5 # Repeat to make big

# --- Main Execution ---

if __name__ == "__main__":
    base_content = get_base_content_large()
    irrelevant_content = get_irrelevant_content()
    
    # 1. Reference File (100% of Base)
    create_docx("1_Reference_File.docx", base_content)
    
    # 2. 85% Similar
    # Mix 85% base with 15% irrelevant (or just slight modifications)
    # To accurately simulate 85% similarity, we keep 85% of the text identical.
    doc_85 = mix_content(base_content, irrelevant_content, 0.85)
    create_docx("2_Highly_Similar_85_Percent.docx", doc_85)
    
    # 3. 70% Similar
    doc_70 = mix_content(base_content, irrelevant_content, 0.70)
    create_docx("3_Partially_Similar_70_Percent.docx", doc_70)
    
    # 4. 30% Similar (Partially Not Similar)
    doc_30 = mix_content(base_content, irrelevant_content, 0.30)
    create_docx("4_Low_Similarity_30_Percent.docx", doc_30)
    
    # 5. 0% Similar
    create_docx("5_No_Similarity_0_Percent.docx", irrelevant_content)
    
    # 6. Structural Variant (Same content, different structure)
    # Ideally this is 100% semantic match but low exact word match.
    # We use the paraphrased generator on the base content.
    structural_variant = generate_paraphrased(base_content)
    create_docx("6_Structural_Variant_Paraphrased.docx", structural_variant)
    
    print("Done generating files.")
