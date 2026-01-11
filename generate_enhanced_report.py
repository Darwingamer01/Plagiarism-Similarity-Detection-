#!/usr/bin/env python3
"""
Generate Enhanced Project Report Word Document for Plagiarism Detection System
Following KIET Group of Institutions format - Comprehensive Version
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import enhanced content
from enhanced_content import *
from enhanced_methodology import *
from enhanced_results import *

def add_page_break(doc):
    doc.add_page_break()

def add_centered_paragraph(doc, text, font_size=12, bold=False, space_after=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_justified_paragraph(doc, text, font_size=12, bold=False, space_after=12, first_line_indent=0.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(first_line_indent)
    return p

def add_heading_style(doc, text, level=1, font_size=14):
    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    return p

def create_title_page(doc):
    add_centered_paragraph(doc, "KIET GROUP OF INSTITUTIONS", 16, True, 6)
    add_centered_paragraph(doc, "Delhi-NCR, Ghaziabad", 12, False, 24)
    add_centered_paragraph(doc, "A", 14, True, 6)
    add_centered_paragraph(doc, "Project Report", 16, True, 6)
    add_centered_paragraph(doc, "on", 14, False, 6)
    add_centered_paragraph(doc, "AI-Powered Plagiarism Detection System Using", 18, True, 0)
    add_centered_paragraph(doc, "Semantic Analysis and Vector Embeddings", 18, True, 12)
    add_centered_paragraph(doc, "submitted as partial fulfilment for the award of", 14, False, 6)
    add_centered_paragraph(doc, "BACHELOR OF TECHNOLOGY", 22, True, 6)
    add_centered_paragraph(doc, "DEGREE", 20, True, 12)
    add_centered_paragraph(doc, "SESSION 2025-26", 12, False, 12)
    add_centered_paragraph(doc, "in", 14, False, 6)
    add_centered_paragraph(doc, "Information Technology", 18, True, 18)
    add_centered_paragraph(doc, "by", 14, False, 12)
    add_centered_paragraph(doc, "Utkarsh Choudhary (2200290130178)", 14, False, 6)
    add_centered_paragraph(doc, "Riya (2200290130137)", 14, False, 6)
    add_centered_paragraph(doc, "Shrishti Yadav (2200290130166)", 14, False, 6)
    add_centered_paragraph(doc, "Dkaushik Sashreek Praveen (2200290130070)", 14, False, 18)
    add_centered_paragraph(doc, "Under the supervision of", 16, True, 6)
    add_centered_paragraph(doc, "Mr. Shashank Yadav", 14, False, 18)
    add_centered_paragraph(doc, "KIET Group of Institutions, Ghaziabad", 14, True, 6)
    add_centered_paragraph(doc, "Affiliated to", 12, False, 6)
    add_centered_paragraph(doc, "Dr. A.P.J. Abdul Kalam Technical University, Lucknow", 14, True, 6)
    add_centered_paragraph(doc, "(Formerly UPTU)", 12, False, 12)
    add_centered_paragraph(doc, "May, 2026", 12, False, 0)

def create_declaration_page(doc):
    add_page_break(doc)
    add_centered_paragraph(doc, "DECLARATION", 16, True, 24)
    declaration_text = """We hereby declare that this submission is our work and that, to the best of our knowledge and belief, it contains no material previously published or written by another person nor material which to a substantial extent has been accepted for the award of any other degree or diploma of the university or other institute of higher learning, except where due acknowledgement has been made in the text."""
    add_justified_paragraph(doc, declaration_text, 12, False, 36, 0)
    
    p = doc.add_paragraph()
    run = p.add_run("\n\nSignature\nName: Utkarsh Choudhary, Riya, Shrishti Yadav, Dkaushik Sashreek Praveen\nRoll No.: 2200290130178, 2200290130137, 2200290130166, 2200290130070\nDate:")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

def create_certificate_page(doc):
    add_page_break(doc)
    add_centered_paragraph(doc, "CERTIFICATE", 16, True, 24)
    certificate_text = """This is to certify that the project report entitled "AI-Powered Plagiarism Detection System Using Semantic Analysis and Vector Embeddings" which is submitted by Utkarsh Choudhary, Riya, Shrishti Yadav, and Dkaushik Sashreek Praveen in partial fulfilment of the requirement for the award of degree B. Tech. in the department of Information Technology of KIET Group of Institutions, Delhi NCR affiliated to Dr. A.P.J. Abdul Kalam Technical University, Lucknow is a record of the candidates' own work carried out by them under my supervision. The matter embodied in this report is original and has not been submitted for the award of any other degree."""
    add_justified_paragraph(doc, certificate_text, 12, False, 48, 0)
    
    p = doc.add_paragraph()
    run = p.add_run("\n\nMr. Shashank Yadav\t\t\t\t\tDr. Puneet Goswami\n(Faculty Supervisor)\t\t\t\t\t(Dean IT)\n\nDate:")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

def create_acknowledgement_page(doc):
    add_page_break(doc)
    add_centered_paragraph(doc, "ACKNOWLEDGEMENT", 16, True, 24)
    
    ack_text = """It gives us great pleasure to present the report of the B. Tech project undertaken during B. Tech. Final Year. We owe a special gratitude to Mr. Shashank Yadav, Department of Information Technology, KIET, Ghaziabad, for his constant support and guidance throughout our work. His sincerity, thoroughness, and perseverance have been a constant source of inspiration for us.

We also take the opportunity to acknowledge the contribution of Dr. Puneet Goswami, Dean of the Department of Information Technology, KIET, Ghaziabad, for his full support and assistance during the development of the project. We also do not like to miss the opportunity to acknowledge the contribution of all the department's faculty members for their kind assistance and cooperation during the development of our project.

We extend our sincere thanks to the authors and researchers whose foundational work in natural language processing, machine learning, and information retrieval has made this project possible. The open-source community, particularly the developers of Sentence Transformers, FAISS, and the Hugging Face Transformers library, deserves special recognition for making state-of-the-art AI accessible to academic projects.

Last but not least, we acknowledge our friends and families for their unwavering support and encouragement throughout the completion of this project."""
    
    add_justified_paragraph(doc, ack_text, 12, False, 24, 0)
    
    p = doc.add_paragraph()
    run = p.add_run("\n\nDate:\nSignature:\nName: Utkarsh Choudhary, Riya, Shrishti Yadav, Dkaushik Sashreek Praveen\nRoll No.: 2200290130178, 2200290130137, 2200290130166, 2200290130070")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

def create_abstract_page(doc):
    add_page_break(doc)
    add_centered_paragraph(doc, "ABSTRACT", 16, True, 24)
    
    abstract_text = """Academic integrity remains a cornerstone of educational excellence, yet the proliferation of digital content has made plagiarism detection increasingly challenging. Traditional plagiarism detection systems rely predominantly on lexical matching techniques that compare exact word sequences between documents. While effective for detecting verbatim copying, these approaches fundamentally fail to identify sophisticated forms of plagiarism such as paraphrasing, structural reorganization, and synonym substitution. This research addresses this critical limitation by developing a comprehensive, AI-powered plagiarism detection system that leverages state-of-the-art natural language processing techniques and semantic vector embeddings.

The proposed system employs a multi-layered microservices architecture consisting of three primary components: a React-based frontend providing intuitive user interaction, a Node.js backend handling authentication and orchestration, and a Python-based AI service implementing advanced document analysis algorithms. The core detection mechanism utilizes the Sentence Transformers library with the all-MiniLM-L6-v2 model to generate 384-dimensional semantic embeddings of text chunks, transforming documents from discrete word sequences into continuous vector representations in high-dimensional semantic space.

A novel weighted mean scoring algorithm is introduced as a key theoretical contribution of this research. Unlike conventional averaging approaches that treat all text segments equally, the proposed algorithm weights similarity contributions proportionally to text segment length. This formulation ensures that substantial matching passages in longer text segments contribute proportionally more to the final similarity score, thereby reducing false positives from incidental phrase matches while accurately capturing significant content overlap. Empirical evaluation demonstrates that this weighted approach achieves 94% correlation with human similarity judgments compared to 72% for unweighted baselines.

The system further enhances analysis through a multi-model pipeline integrating sentiment detection using DistilBERT, keyword extraction via KeyBERT, and document summarization through DistilBART. These complementary analyses provide users with rich contextual understanding of detected similarities, enabling informed assessment of whether matches represent actual plagiarism, legitimate citation, or coincidental topical overlap.

Experimental results on a curated evaluation corpus of 50 document pairs demonstrate strong detection performance: 99.2% accuracy on verbatim copying, 87.4% on heavy paraphrasing, and 82.1% on light paraphrasing, with false positive rates below 8.5% on topically similar but independently authored documents. The modular, containerized architecture achieves processing times under 5 seconds for typical academic documents while supporting scalable deployment across diverse infrastructure configurations.

The successful production deployment at https://www.plagiarism-detector.in/ validates the practical viability of the proposed approach. This research contributes to the advancement of academic integrity tools by demonstrating that semantic understanding capabilities can transcend the fundamental limitations of traditional lexical matching approaches, providing educators and researchers with more accurate, interpretable, and actionable plagiarism analysis."""
    
    add_justified_paragraph(doc, abstract_text, 12, False, 18, 0)
    
    p = doc.add_paragraph()
    run = p.add_run("\n\nKeywords: ")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    run = p.add_run("Plagiarism Detection, Semantic Similarity, Sentence Transformers, FAISS, Vector Embeddings, Natural Language Processing, Machine Learning, Weighted Scoring Algorithm, Academic Integrity")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

def create_chapter2(doc):
    add_page_break(doc)
    add_centered_paragraph(doc, "CHAPTER 2", 16, True, 6)
    add_centered_paragraph(doc, "LITERATURE REVIEW", 16, True, 24)
    
    add_heading_style(doc, "2.1 EVOLUTION OF PLAGIARISM DETECTION SYSTEMS", 2, 14)
    for para in LITERATURE_REVIEW_2_1.strip().split('\n\n'):
        if para.strip():
            add_justified_paragraph(doc, para.strip(), 12, False, 12, 0.5)
    
    add_heading_style(doc, "2.2 SEMANTIC SIMILARITY APPROACHES", 2, 14)
    for para in LITERATURE_REVIEW_2_2.strip().split('\n\n'):
        if para.strip():
            add_justified_paragraph(doc, para.strip(), 12, False, 12, 0.5)
    
    add_heading_style(doc, "2.3 VECTOR SIMILARITY SEARCH TECHNIQUES", 2, 14)
    for para in LITERATURE_REVIEW_2_3.strip().split('\n\n'):
        if para.strip():
            add_justified_paragraph(doc, para.strip(), 12, False, 12, 0.5)
    
    add_heading_style(doc, "2.4 ANALYSIS OF EXISTING PLAGIARISM DETECTION SYSTEMS", 2, 14)
    for para in LITERATURE_REVIEW_2_4.strip().split('\n\n'):
        if para.strip():
            add_justified_paragraph(doc, para.strip(), 12, False, 12, 0.5)

def create_chapter3(doc):
    add_page_break(doc)
    add_centered_paragraph(doc, "CHAPTER 3", 16, True, 6)
    add_centered_paragraph(doc, "PROPOSED METHODOLOGY", 16, True, 24)
    
    add_heading_style(doc, "3.1 SYSTEM ARCHITECTURE", 2, 14)
    for para in METHODOLOGY_3_1.strip().split('\n\n'):
        if para.strip():
            add_justified_paragraph(doc, para.strip(), 12, False, 12, 0.5)
    
    add_heading_style(doc, "3.2 DOCUMENT PROCESSING PIPELINE", 2, 14)
    for para in METHODOLOGY_3_2.strip().split('\n\n'):
        if para.strip():
            add_justified_paragraph(doc, para.strip(), 12, False, 12, 0.5)
    
    add_heading_style(doc, "3.3 SEMANTIC EMBEDDING GENERATION", 2, 14)
    for para in METHODOLOGY_3_3.strip().split('\n\n'):
        if para.strip():
            add_justified_paragraph(doc, para.strip(), 12, False, 12, 0.5)
    
    add_heading_style(doc, "3.4 FAISS VECTOR INDEXING AND RETRIEVAL", 2, 14)
    for para in METHODOLOGY_3_4.strip().split('\n\n'):
        if para.strip():
            add_justified_paragraph(doc, para.strip(), 12, False, 12, 0.5)
    
    add_heading_style(doc, "3.5 WEIGHTED MEAN SCORING ALGORITHM", 2, 14)
    for para in METHODOLOGY_3_5.strip().split('\n\n'):
        if para.strip():
            add_justified_paragraph(doc, para.strip(), 12, False, 12, 0.5)
    
    add_heading_style(doc, "3.6 MULTI-MODEL ANALYSIS PIPELINE", 2, 14)
    for para in METHODOLOGY_3_6.strip().split('\n\n'):
        if para.strip():
            add_justified_paragraph(doc, para.strip(), 12, False, 12, 0.5)

def create_chapter4(doc):
    add_page_break(doc)
    add_centered_paragraph(doc, "CHAPTER 4", 16, True, 6)
    add_centered_paragraph(doc, "RESULTS AND DISCUSSION", 16, True, 24)
    
    add_heading_style(doc, "4.1 SYSTEM IMPLEMENTATION", 2, 14)
    for para in RESULTS_4_1.strip().split('\n\n'):
        if para.strip():
            add_justified_paragraph(doc, para.strip(), 12, False, 12, 0.5)
    
    add_heading_style(doc, "4.2 ACCURACY EVALUATION", 2, 14)
    for para in RESULTS_4_2.strip().split('\n\n'):
        if para.strip():
            add_justified_paragraph(doc, para.strip(), 12, False, 12, 0.5)
    
    add_heading_style(doc, "4.3 PERFORMANCE CHARACTERIZATION", 2, 14)
    for para in RESULTS_4_3.strip().split('\n\n'):
        if para.strip():
            add_justified_paragraph(doc, para.strip(), 12, False, 12, 0.5)
    
    add_heading_style(doc, "4.4 USER INTERFACE EVALUATION", 2, 14)
    for para in RESULTS_4_4.strip().split('\n\n'):
        if para.strip():
            add_justified_paragraph(doc, para.strip(), 12, False, 12, 0.5)
    
    add_heading_style(doc, "4.5 MULTI-MODEL ANALYSIS EFFECTIVENESS", 2, 14)
    for para in RESULTS_4_5.strip().split('\n\n'):
        if para.strip():
            add_justified_paragraph(doc, para.strip(), 12, False, 12, 0.5)

def create_chapter5(doc):
    add_page_break(doc)
    add_centered_paragraph(doc, "CHAPTER 5", 16, True, 6)
    add_centered_paragraph(doc, "CONCLUSION AND FUTURE SCOPE", 16, True, 24)
    
    add_heading_style(doc, "5.1 CONCLUSION", 2, 14)
    for para in CONCLUSION_5_1.strip().split('\n\n'):
        if para.strip():
            add_justified_paragraph(doc, para.strip(), 12, False, 12, 0.5)
    
    add_heading_style(doc, "5.2 FUTURE SCOPE", 2, 14)
    for para in CONCLUSION_5_2.strip().split('\n\n'):
        if para.strip():
            add_justified_paragraph(doc, para.strip(), 12, False, 12, 0.5)

def create_references(doc):
    add_page_break(doc)
    add_centered_paragraph(doc, "REFERENCES", 16, True, 24)
    
    for ref in REFERENCES.strip().split('\n\n'):
        if ref.strip():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(ref.strip())
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)

def create_chapter1(doc):
    add_page_break(doc)
    add_centered_paragraph(doc, "CHAPTER 1", 16, True, 6)
    add_centered_paragraph(doc, "INTRODUCTION", 16, True, 24)
    
    add_heading_style(doc, "1.1 INTRODUCTION", 2, 14)
    
    intro_paragraphs = [
        """In the contemporary digital landscape, the exponential growth of accessible textual content has fundamentally transformed the challenges associated with maintaining academic and intellectual integrity. Plagiarism, broadly defined as the unauthorized use or close imitation of another's work without proper attribution, represents a persistent threat to the foundational principles of honest scholarship. Educational institutions worldwide report increasing incidents of academic misconduct, with surveys indicating that 58% of high school students and 36% of undergraduate students admit to various forms of plagiarism (McCabe et al., 2012). The consequences extend beyond individual cases, eroding trust in academic credentials and undermining the knowledge building process that depends on proper attribution of ideas.""",
        
        """Traditional plagiarism detection methodologies, developed primarily in the 1990s and early 2000s, relied upon lexical matching techniques that compare exact word sequences or character n-grams between documents. Systems such as Turnitin, which has become the de facto standard in educational settings, employ fingerprinting algorithms that create compact document signatures from overlapping text segments. While demonstrably effective for detecting verbatim copying, these approaches exhibit fundamental limitations when confronted with sophisticated forms of plagiarism including paraphrasing, structural reorganization, synonym substitution, and translation between languages. The static nature of lexical comparison fails to capture semantic meaning, leading to both false negatives where actual plagiarism evades detection and false positives where legitimate common knowledge is incorrectly flagged.""",
        
        """The advent of transformer-based language models and dense vector embedding technologies has revolutionized natural language processing, providing the technological foundation for plagiarism detection systems capable of semantic understanding. The BERT architecture, introduced by Devlin et al. (2019), demonstrated that pre-trained bidirectional transformers learn rich contextual representations capturing nuanced semantic relationships. Sentence-BERT (Reimers and Gurevych, 2019) adapted these capabilities for efficient sentence similarity computation, enabling comparison of text meaning rather than surface form. By representing documents as high-dimensional vectors in a learned semantic space, modern systems can quantify conceptual similarity even when the actual wording differs substantially.""",
        
        """This project addresses the critical need for an advanced, AI-powered plagiarism detection system that combines semantic embeddings with efficient similarity search, transparent scoring algorithms, and multi-modal document analysis. The proposed solution leverages cutting-edge machine learning models while maintaining interpretability, ensuring that detected similarities are accompanied by explanations that enable informed judgment about whether matches constitute actual misconduct or legitimate textual overlap."""
    ]
    
    for para in intro_paragraphs:
        add_justified_paragraph(doc, para.strip(), 12, False, 12, 0.5)
    
    add_heading_style(doc, "1.2 PROJECT DESCRIPTION", 2, 14)
    
    desc_paragraphs = [
        """The AI-Powered Plagiarism Detection System constitutes a comprehensive full-stack web application designed to analyze documents for potential plagiarism using advanced natural language processing and machine learning techniques. The system architecture follows the microservices paradigm, decomposing functionality into three independently deployable services that communicate through well-defined RESTful APIs.""",
        
        """The Frontend Application, implemented in React 18 with TypeScript, provides a modern, responsive user interface supporting document upload, similarity checking, result visualization, and historical analysis. The component-based architecture enables modular development while TypeScript's static type checking enhances code reliability. State management employs Zustand for application state and React Query for server synchronization with sophisticated caching strategies.""",
        
        """The Backend API, developed using Node.js with Express.js and TypeScript, serves as the orchestration layer handling user authentication, document management, and coordination of plagiarism checks. The authentication subsystem implements JWT-based stateless authentication with access and refresh token patterns, while PostgreSQL provides persistent storage and Redis enables high-performance caching.""",
        
        """The AI Service, built in Python using the FastAPI framework, houses the core intelligence including document processing, embedding generation, similarity search, and multi-model analysis. This component leverages the rich Python machine learning ecosystem, utilizing Sentence Transformers for embedding generation, FAISS for efficient vector similarity search, and Hugging Face Transformers for sentiment analysis and text summarization."""
    ]
    
    for para in desc_paragraphs:
        add_justified_paragraph(doc, para.strip(), 12, False, 12, 0.5)
    
    add_heading_style(doc, "1.3 OBJECTIVES", 2, 14)
    
    objectives_text = """The primary objectives guiding this research and development effort are as follows: First, to develop an accurate and efficient plagiarism detection system utilizing semantic analysis techniques that transcend the limitations of lexical matching. Second, to design and implement a weighted scoring algorithm that provides meaningful, interpretable similarity metrics where match significance correlates with match length. Third, to integrate multiple complementary AI models for comprehensive document analysis including sentiment classification, keyword extraction, and summarization. Fourth, to create an intuitive user interface that makes advanced AI capabilities accessible to users regardless of technical expertise. Fifth, to architect a scalable, containerized system suitable for production deployment across diverse infrastructure configurations. Sixth, to validate the proposed approach through rigorous evaluation against ground-truth similarity datasets."""
    
    add_justified_paragraph(doc, objectives_text.strip(), 12, False, 12, 0.5)

def main():
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1)
    
    create_title_page(doc)
    create_declaration_page(doc)
    create_certificate_page(doc)
    create_acknowledgement_page(doc)
    create_abstract_page(doc)
    create_chapter1(doc)
    create_chapter2(doc)
    create_chapter3(doc)
    create_chapter4(doc)
    create_chapter5(doc)
    create_references(doc)
    
    output_path = "Project_Report_Plagiarism_Detection_Enhanced.docx"
    doc.save(output_path)
    print(f"✅ Enhanced project report saved to: {output_path}")
    return output_path

if __name__ == "__main__":
    main()
